from pathlib import Path
from datetime import datetime
from typing import Dict

from pypdf import PdfReader
from docx import Document


def _base_result(path: str, file_type: str, title: str, content: str) -> Dict[str, str]:
    return {
        "source_path": str(Path(path).resolve()),
        "file_type": file_type,
        "title": title,
        "content": content,
        "loaded_at": datetime.utcnow().isoformat(),
    }


def detect_file_type(path: str) -> str:
    p = Path(path)
    ext = p.suffix.lower()
    if ext in {".txt", ".md", ".pdf", ".docx"}:
        return ext
    raise ValueError(f"Unsupported file type: {ext}")

#处理 .txt .md 文件
def load_text(path: str) -> Dict[str, str]:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # TODO: 读取文本内容（utf-8）
    # 建议：加 errors="ignore" 防止编码问题
    content = p.read_text(encoding="utf-8", errors= "ignore")

    if not content.strip():
        print(f"[WARN] Empty text content: {path}")

    return _base_result(
        path=str(p),
        file_type=p.suffix.lower(),
        title=p.stem,
        content=content,
    )


def load_pdf(path: str) -> Dict[str, str]:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # TODO: PdfReader + 遍历 pages 提取文本
    # content = "\n".join(...)
    reader = PdfReader(str(p))
    parts = []

    for i,page in enumerate(reader.pages,start = 1):
        text  = page.extract_text() or ""
        if text.strip():
            parts.append(text)
        else:
            print(f"[WARN] PDF page {i} has no extractable text: {path}")

    content = "\n".join(parts).strip()

    if not content:
        print(f"[WARN] Empty pdf content: {path}")

    return _base_result(
        path=str(p),
        file_type=".pdf",
        title=p.stem,
        content=content,
    )


def load_docx(path: str) -> Dict[str, str]:
    p = Path(path)
    doc = Document(str(p))
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # TODO: Document(path) + 读取 paragraphs
    content = "\n".join(p.text for p in doc.paragraphs)

    if not content:
        print(f"[WARN] Empty docx content: {path}")

    return _base_result(
        path=str(p),
        file_type=".docx",
        title=p.stem,
        content=content,
    )


def load_document(path: str) -> Dict[str, str]:
    file_type = detect_file_type(path)

    if file_type in {".txt", ".md"}:
        return load_text(path)
    if file_type == ".pdf":
        return load_pdf(path)
    if file_type == ".docx":
        return load_docx(path)

    # 理论上不会到这里（detect_file_type 已拦截）
    raise ValueError(f"Unsupported file type: {file_type}")
